import io
import docx
import json
from typing import Dict, Any, Optional
from pypdf import PdfReader
import pytesseract
from PIL import Image

class TextExtractor:
    """
    Service responsible for extracting text from various file formats.

    This service can operate as part of a batch ingestion workflow or
    process files directly from memory. Supported formats include
    PDF, DOCX, JSON, plain text, and common image types (JPEG, PNG).

    Extraction methods are pluggable and mapped by file extension.
    """

    def __init__(self, storage_client=None):
        """
        Initialize the text extractor.

        Args:
            storage_client: Optional storage client to load files from a
                remote or local storage system if file bytes are not
                directly provided.
        """
        self.storage = storage_client
        self.extraction_methods = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.json': self._extract_from_json,
            '.txt': self._extract_from_plaintext,
            '.text': self._extract_from_plaintext,
            '.jpg': self._extract_from_image,
            '.jpeg': self._extract_from_image,
            '.png': self._extract_from_image,
        }

    def extract_text(
            self,
            file_path: str,
            metadata: Dict[str, Any],
            file_content: Optional[bytes] = None,
            extension: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract text from a file and return structured data.

        Workflow:
            1. Determine content source (in-memory bytes or storage client)
            2. Determine file extension
            3. Select the appropriate extraction method
            4. Extract text
            5. Return JSON with extracted text and metadata

        Args:
            file_path: Path or name of the file to process
            metadata: Initial metadata dictionary (e.g., user_id, doc_id)
            file_content: Optional file content as bytes
            extension: Optional file extension (automatically detected if missing)

        Returns:
            Dict containing:
                - doc_id: Document identifier from metadata
                - file_path: Original file path
                - file_type: File extension
                - extracted_text: Extracted text content
                - metadata: Original metadata dictionary
                - error: Optional error message if extraction failed
        """
        try:
            if file_content is not None:
                content = file_content
                file_ext = extension or self._get_extension_from_path(file_path)
            elif self.storage:
                content = self.storage.load_file_content(file_path)
                file_ext = self.storage.get_file_extension(file_path)
            else:
                return {
                    "error": "No content provided and no storage client available",
                    "file_path": file_path
                }

            extractor = self.extraction_methods.get(file_ext)
            if extractor:
                extracted_text = extractor(content)
            else:
                return {
                    "error": f"Unsupported or unknown file format: {file_ext}",
                    "file_path": file_path,
                    "file_type": file_ext
                }

            result_json = {
                "doc_id": metadata.get("doc_id"),
                "file_path": file_path,
                "file_type": file_ext,
                "extracted_text": extracted_text,
                "metadata": metadata
            }
            return result_json

        except FileNotFoundError as e:
            return {"error": f"File not found: {e}", "file_path": file_path}
        except Exception as e:
            return {"error": f"Extraction failed: {e}", "file_path": file_path}

    @staticmethod
    def _get_extension_from_path(file_path: str) -> str:
        """
        Extract file extension from a file path.

        Args:
            file_path: Path or filename

        Returns:
            Normalized lowercase extension (e.g., '.pdf'), or empty string if none
        """
        if '.' in file_path:
            return '.' + file_path.split('.')[-1].lower()
        return ''

    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """
        Extract text from a PDF file.

        Args:
            content: File content as bytes

        Returns:
            Extracted text as a string
        """
        reader = PdfReader(io.BytesIO(content))

        if len(reader.pages) == 0:
            return "Error: PDF contains no pages"

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)

        if not text_parts:
            return "Warning: No text could be extracted from PDF (might be image-based)"

        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """
        Extract text from a DOCX file.

        Args:
            content: File content as bytes

        Returns:
            Extracted text as a string
        """
        doc = docx.Document(io.BytesIO(content))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        all_text = []
        if paragraphs:
            all_text.append('\n'.join(paragraphs))
        if table_texts:
            all_text.append('--- Tables ---\n' + '\n'.join(table_texts))

        if not all_text:
            return "Warning: No text could be extracted from DOCX"

        return '\n\n'.join(all_text)

    @staticmethod
    def _extract_from_json(content: bytes) -> str:
        """
        Extract relevant text from a JSON file.

        Tries to read a 'text' field, falling back to serialized JSON.

        Args:
            content: File content as bytes

        Returns:
            Extracted or serialized JSON as string
        """
        try:
            data = json.loads(content.decode('utf-8'))
            return data.get('text', json.dumps(data))
        except json.JSONDecodeError:
            return content.decode('utf-8')

    @staticmethod
    def _extract_from_plaintext(content: bytes) -> str:
        """
        Decode plaintext content.

        Tries UTF-8 first, falls back to Latin-1 ignoring errors.

        Args:
            content: File content as bytes

        Returns:
            Decoded text string
        """
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('latin-1', errors='ignore')

    @staticmethod
    def _extract_from_image(content: bytes) -> str:
        """
        Extract text from an image using OCR (Optical Character Recognition).

        Args:
            content: Image content as bytes

        Returns:
            Extracted text as a string
        """
        image = Image.open(io.BytesIO(content))

        text = pytesseract.image_to_string(image, lang='eng')

        text = text.strip()

        if not text:
            return "Warning: No text could be extracted from image (might be blank or poor quality)"

        return text